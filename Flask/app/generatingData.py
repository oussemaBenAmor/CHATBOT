import os
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib import colors
from reportlab.lib.units import inch
from docx import Document
from docx.shared import Inches
import openpyxl
from datetime import datetime

# Define transaction types, condition variations, section titles, and explanations
TRANSACTION_TYPES = {
    "refunds": {
        "section_titles": ["Refund Eligibility", "Return Requirements", "Refund Procedures"],
        "versions": [
            # ...existing versions...
            [
                ("Refunds are issued within 14 days for defective products.", 
                 "Only products verified as defective by our team qualify for refunds within this period."),
                ("No refunds are provided for items purchased more than 30 days ago.", 
                 "This policy ensures timely processing and verification of refund requests."),
                ("Damaged items must be returned in original packaging for a refund.", 
                 "Original packaging helps verify the condition of the returned item."),
                ("Credit card refunds are processed to the original card within 5 business days.", 
                 "Refunds are credited directly to the card used for purchase."),
                ("Refunds are not available for change of mind on non-defective items.", 
                 "This applies to items that are fully functional and undamaged.")
            ],
            [
                ("Refunds are available within 10 days for faulty items.", 
                 "Faulty items must be reported and returned within this timeframe."),
                ("Non-defective items are not refundable after 14 days.", 
                 "This ensures returns are processed promptly after purchase."),
                ("Refunds for damaged goods require proof of damage via email.", 
                 "Submit photos to support@company.com to initiate the refund process."),
                ("Refunds to bank accounts take 7 business days to process.", 
                 "Bank account details must be provided for non-card refunds."),
                ("No refunds for international orders unless defective.", 
                 "International orders are subject to stricter refund policies.")
            ],
            [
                ("Refunds are processed within 7 days for eligible products.", 
                 "Eligibility is determined by product condition and purchase date."),
                ("Items must be unused and in saleable condition for a refund.", 
                 "Saleable condition means no signs of use or damage."),
                ("Refunds for incorrect orders are issued after return verification.", 
                 "Verification ensures the correct item is returned."),
                ("Credit card refunds are credited within 3 business days.", 
                 "Refunds are processed to the original payment method swiftly."),
                ("No refunds for promotional items or special orders.", 
                 "Promotional items are final sale due to discounted pricing.")
            ],
            [
                ("All refunds require a valid receipt or order number.", 
                 "This ensures eligibility for all refund processing."),  # Common condition
                ("Refunds for perishable goods are limited to 3 days.", 
                 "Perishable items must be returned quickly due to spoilage risks."),
                ("Partial refunds are available for partially used items.", 
                 "A prorated refund is calculated based on usage."),
                ("Refunds can be requested via our mobile app.", 
                 "The app provides a streamlined refund submission process."),
                ("No refunds are issued for items lost in transit.", 
                 "Lost items are covered under shipping insurance instead.")
            ],
            [
                ("All refunds require a valid receipt or order number.", 
                 "This ensures eligibility for all refund processing."),  # Common condition
                ("Refunds for seasonal items are only valid until stock clearance.", 
                 "Seasonal items are non-returnable after clearance sales."),
                ("Refunds include a 2% processing fee for expedited requests.", 
                 "Expedited refunds are processed within 24 hours."),
                ("Refunds for bulk orders require manager approval.", 
                 "Approval ensures compliance with bulk purchase terms."),
                ("Refunds are paused during system maintenance periods.", 
                 "Maintenance schedules are posted on our website.")
            ],
            [
                ("All refunds require a valid receipt or order number.", 
                 "This ensures eligibility for all refund processing."),  # Common condition
                ("Refunds for custom orders are not permitted.", 
                 "Custom items are made to order and non-returnable."),
                ("Refunds are available in store credit for first-time returns.", 
                 "Credit can be used for future purchases."),
                ("Refunds for recalled items are processed within 1 day.", 
                 "Recall refunds prioritize customer safety."),
                ("Refunds are denied if packaging is missing labels.", 
                 "Labels are required for return identification.")
            ],
            [
                ("All refunds require a valid receipt or order number.", 
                 "This ensures eligibility for all refund processing."),  # Common condition
                ("Refunds for high-value items require a quality inspection.", 
                 "Inspection verifies the item’s condition before refund."),
                ("Refunds can be converted to gift cards at a 5% bonus.", 
                 "Bonus incentivizes use of gift cards for future purchases."),
                ("Refunds for defective batches are processed in bulk.", 
                 "Bulk processing speeds up refunds for affected customers."),
                ("No refunds are available during product launches.", 
                 "Launches prioritize new product distribution.")
            ],
            [
                ("All refunds require a valid receipt or order number.", 
                 "This ensures eligibility for all refund processing."),  # Common condition
                ("Refunds for returned gifts require donor consent.", 
                 "Consent ensures the original purchaser agrees to the return."),
                ("Refunds are issued as cash for returns over $100.", 
                 "Cash refunds apply to higher-value transactions."),
                ("Refunds for warranty claims are expedited by 2 days.", 
                 "Expedited processing supports warranty-covered items."),
                ("Refunds are tracked via a unique refund ID.", 
                 "ID allows customers to monitor refund status online.")
            ],
            [
                ("All refunds require a valid receipt or order number.", 
                 "This ensures eligibility for all refund processing."),  # Common condition
                ("Refunds for eco-friendly products support recycling fees.", 
                 "Fees contribute to sustainable disposal programs."),
                ("Refunds are denied if items are altered post-purchase.", 
                 "Alterations void refund eligibility."),
                ("Refunds for subscription cancellations are prorated.", 
                 "Proration reflects the unused subscription period."),
                ("Refunds are processed offline during network failures.", 
                 "Offline processing ensures continuity during outages.")
            ],
            # --- New version with cross-transaction relationships ---
            [
                # Refunds related to payments
                ("Payments made by credit card are refundable if the item is unused and in saleable condition.",
                 "Refunds for credit card payments are processed within 5 business days to the original card."),
                ("Payments made by bank transfer are refundable to the original bank account within 7 business days.",
                 "A valid bank account must be provided for refund processing."),
                ("Refunds are not available for payments made in cash for international orders.",
                 "International cash payments are final and non-refundable."),
                # Refunds related to transfers
                ("Failed transfers are automatically refunded within 48 hours.",
                 "Refunds for failed transfers are processed to the original payment method."),
                # Refunds related to exchanges
                ("If an exchange cannot be fulfilled due to out-of-stock items, a refund will be issued.",
                 "Refunds for unfulfilled exchanges follow the standard refund process."),
                # General
                ("Refund requests for payments made during promotional events are subject to special conditions.",
                 "Promotional event refunds may be limited or delayed."),
                ("Refunds for items received via exchange are only available if the exchanged item is defective.",
                 "Refunds are not issued for change of mind after an exchange."),
                ("Refunds for payments made by cryptocurrency are processed at the current exchange rate.",
                 "Cryptocurrency refunds may be subject to conversion fees."),
                ("Refunds for bulk payments require manager approval and may be prorated.",
                 "Approval ensures compliance with bulk purchase terms."),
                ("Refunds for transfers between linked accounts are processed within 24 hours.",
                 "Linked account refunds are expedited for customer convenience.")
            ]
        ]
    },
    "payments": {
        "section_titles": ["Payment Processing", "Payment Methods", "Payment Rules"],
        "versions": [
            # ...existing versions...
            [
                ("Payments by credit card are processed immediately.", 
                 "Credit card transactions are verified and completed instantly."),
                ("Bank transfers require a valid account number and take 3 business days.", 
                 "Valid account details must be provided to avoid delays."),
                ("Payments via cheque must clear before order fulfillment.", 
                 "Cheque clearance ensures funds are available for processing."),
                ("Refunds for payments made by bank transfer require a nominated bank account.", 
                 "A valid bank account is needed for refund processing."),
                ("All payment disputes must be reported within 7 days.", 
                 "Timely reporting ensures quick resolution of payment issues.")
            ],
            [
                ("Credit card payments are verified instantly.", 
                 "Verification ensures secure and immediate processing."),
                ("Bank transfers incur a 1% processing fee for international accounts.", 
                 "This fee covers additional banking costs for international transfers."),
                ("Cheque payments require 5 business days to clear.", 
                 "Clearance time ensures funds are available before shipping."),
                ("Payment refunds are processed to the original method within 5 days.", 
                 "Refunds are returned to the original payment source for consistency."),
                ("Disputes for payments must be filed within 10 days.", 
                 "Extended reporting period allows flexibility for dispute resolution.")
            ],
            [
                ("Payments via debit card are processed within 24 hours.", 
                 "Debit card transactions are processed quickly for efficiency."),
                ("International bank transfers take 5 business days to complete.", 
                 "Additional time is needed for international banking protocols."),
                ("Payments by cash are accepted for in-store purchases only.", 
                 "Cash payments are limited to physical store transactions."),
                ("Refunds for payments are issued to the original account within 7 days.", 
                 "Original account refunds ensure traceability and security."),
                ("Payment issues must be reported within 3 business days.", 
                 "Prompt reporting helps resolve issues swiftly.")
            ],
            [
                ("All payments must include a transaction reference number.", 
                 "This ensures accurate tracking of each payment."),  # Common condition
                ("Payments via mobile wallet are processed within 2 hours.", 
                 "Mobile wallet transactions use real-time verification."),
                ("Split payments are allowed with a 1.5% surcharge.", 
                 "Surcharge covers additional processing for split transactions."),
                ("Payments exceeding $500 require two-factor authentication.", 
                 "Security is enhanced for large transactions."),
                ("Payment failures trigger an automatic retry after 24 hours.", 
                 "Retries ensure payment completion without manual intervention.")
            ],
            [
                ("All payments must include a transaction reference number.", 
                 "This ensures accurate tracking of each payment."),  # Common condition
                ("Payments by cryptocurrency are accepted with a 3% fee.", 
                 "Cryptocurrency payments are converted to local currency."),
                ("Recurring payments require annual renewal approval.", 
                 "Approval prevents unauthorized recurring charges."),
                ("Payments made during holidays are delayed by 1 day.", 
                 "Delays account for reduced banking operations."),
                ("Payment disputes can be escalated to a supervisor.", 
                 "Escalation provides a higher level of review.")
            ],
            [
                ("All payments must include a transaction reference number.", 
                 "This ensures accurate tracking of each payment."),  # Common condition
                ("Payments via prepaid cards have a $50 minimum.", 
                 "Minimum ensures cost-effective processing."),
                ("Bulk payments receive a 2% discount if paid early.", 
                 "Early payment incentivizes prompt settlement."),
                ("Payments rejected due to fraud are flagged for review.", 
                 "Review prevents future fraudulent attempts."),
                ("Payment processing stops during server outages.", 
                 "Outages are announced via email notifications.")
            ],
            [
                ("All payments must include a transaction reference number.", 
                 "This ensures accurate tracking of each payment."),  # Common condition
                ("Payments via smartwatch are processed within 1 hour.", 
                 "Smartwatch payments use advanced authentication."),
                ("Payments under $10 incur a $0.50 convenience fee.", 
                 "Fee covers small transaction processing costs."),
                ("Multiple payment methods can be combined per order.", 
                 "Combination allows flexibility for large purchases."),
                ("Payment confirmations are sent via SMS.", 
                 "SMS ensures immediate customer notification.")
            ],
            [
                ("All payments must include a transaction reference number.", 
                 "This ensures accurate tracking of each payment."),  # Common condition
                ("Payments by wire transfer require a 48-hour notice.", 
                 "Notice allows time for bank coordination."),
                ("Payments over $1,000 receive a loyalty discount.", 
                 "Discount rewards high-value customers."),
                ("Payment errors are corrected within 72 hours.", 
                 "Corrections ensure accurate billing."),
                ("Payments are paused during security audits.", 
                 "Audits enhance system integrity.")
            ],
            [
                ("All payments must include a transaction reference number.", 
                 "This ensures accurate tracking of each payment."),  # Common condition
                ("Payments via installment plans require credit approval.", 
                 "Approval assesses customer creditworthiness."),
                ("Emergency payments are processed within 30 minutes.", 
                 "Fast processing supports urgent needs."),
                ("Payments made in foreign currency incur a 2% conversion fee.", 
                 "Fee covers currency exchange costs."),
                ("Payment records are archived for 5 years.", 
                 "Archiving complies with regulatory requirements.")
            ],
            # --- New version with cross-transaction relationships ---
            [
                ("Payments made by credit card are eligible for refunds if the item meets refund conditions.",
                 "Refunds for credit card payments are processed within 5 business days."),
                ("Payments made by bank transfer can be refunded to the original account within 7 business days.",
                 "A valid bank account is required for refund processing."),
                ("Payments for exchanged items are adjusted based on price differences.",
                 "If the exchanged item is cheaper, the difference is refunded; if more expensive, additional payment is required."),
                ("Payments for failed transfers are refunded automatically within 48 hours.",
                 "Refunds for failed transfers are processed to the original payment method."),
                ("Payments made during promotional events may have limited refund options.",
                 "Check the promotional terms for refund eligibility."),
                ("Payments for bulk orders are refundable only with manager approval.",
                 "Bulk order refunds may be prorated based on usage."),
                ("Payments made by cryptocurrency are refundable at the current exchange rate.",
                 "Conversion fees may apply for cryptocurrency refunds."),
                ("Payments for international orders are non-refundable if paid in cash.",
                 "International cash payments are final."),
                ("Payments for items received via exchange are subject to the exchange policy.",
                 "Refunds for exchanged items are only available if defective."),
                ("Payments between linked accounts can be reversed within 24 hours if a refund is requested.",
                 "Reversals are expedited for linked accounts.")
            ]
        ]
    },
    "transfers": {
        "section_titles": ["Transfer Guidelines", "Transfer Policies", "Transfer Conditions"],
        "versions": [
            # ...existing versions...
            [
                ("Transfers are processed within 3 business days for domestic accounts.", 
                 "Domestic transfers are prioritized for quick processing."),
                ("International transfers incur a 2% processing fee.", 
                 "This fee covers international banking charges."),
                ("Transfers require verification of recipient bank details.", 
                 "Verification ensures accuracy and security of transfers."),
                ("No transfers are allowed to unverified accounts.", 
                 "Unverified accounts pose a risk to transaction security."),
                ("Transfer disputes must be reported within 5 days.", 
                 "Timely reporting ensures disputes are handled efficiently.")
            ],
            [
                ("Domestic transfers are completed within 2 business days.", 
                 "Faster processing for domestic accounts enhances efficiency."),
                ("International transfers require a 3% fee and 5-day processing.", 
                 "Higher fees and time account for international regulations."),
                ("Recipient details must be verified via email before transfer.", 
                 "Email verification adds an extra layer of security."),
                ("Transfers to unverified accounts are not permitted.", 
                 "This policy prevents fraudulent transactions."),
                ("Disputes for transfers must be reported within 7 days.", 
                 "Extended reporting period allows for thorough dispute resolution.")
            ],
            [
                ("Transfers within the same bank are processed within 24 hours.", 
                 "Same-bank transfers are expedited for customer convenience."),
                ("International transfers incur a 1.5% fee and take 4 business days.", 
                 "Reduced fees and time improve international transfer efficiency."),
                ("Verification of recipient details is mandatory for all transfers.", 
                 "Mandatory verification ensures secure transactions."),
                ("No transfers are processed without a valid recipient account.", 
                 "Valid accounts are required to complete transfers."),
                ("Transfer disputes must be filed within 3 days.", 
                 "Quick reporting ensures rapid dispute resolution.")
            ],
            [
                ("All transfers must include a unique transaction ID.", 
                 "This ensures traceability for all transfer operations."),  # Common condition
                ("Transfers to new accounts require a 24-hour hold.", 
                 "Hold period verifies new account legitimacy."),
                ("Large transfers over $1,000 incur a 1% surcharge.", 
                 "Surcharge covers additional security checks."),
                ("Transfers can be scheduled for future dates.", 
                 "Scheduling allows flexibility for planned transfers."),
                ("Transfer limits are reduced during peak seasons.", 
                 "Limits prevent system overload during high demand.")
            ],
            [
                ("All transfers must include a unique transaction ID.", 
                 "This ensures traceability for all transfer operations."),  # Common condition
                ("International transfers require a compliance review.", 
                 "Review ensures adherence to global regulations."),
                ("Transfers between linked accounts are fee-free.", 
                 "Linked accounts benefit from reduced costs."),
                ("Failed transfers are refunded within 48 hours.", 
                 "Refunds are processed automatically after failure."),
                ("Transfers are paused during banking holidays.", 
                 "Pauses align with official banking schedules.")
            ],
            [
                ("All transfers must include a unique transaction ID.", 
                 "This ensures traceability for all transfer operations."),  # Common condition
                ("Transfers to charity accounts receive a tax receipt.", 
                 "Receipts are issued for tax deduction purposes."),
                ("Recurring transfers require monthly confirmation.", 
                 "Confirmation prevents unauthorized repeats."),
                ("Transfers exceeding $5,000 need executive approval.", 
                 "Approval ensures compliance with high-value policies."),
                ("Transfer history is available for 90 days online.", 
                 "Online access aids in tracking past transactions.")
            ],
            [
                ("All transfers must include a unique transaction ID.", 
                 "This ensures traceability for all transfer operations."),  # Common condition
                ("Transfers to joint accounts require both signatures.", 
                 "Signatures ensure authorization from all parties."),
                ("Instant transfers are available for a 2% fee.", 
                 "Fee covers expedited processing costs."),
                ("Transfers to savings accounts are limited to once monthly.", 
                 "Limit prevents frequent account disruptions."),
                ("Transfer requests are canceled after 24 hours if unverified.", 
                 "Unverified requests are automatically voided.")
            ],
            [
                ("All transfers must include a unique transaction ID.", 
                 "This ensures traceability for all transfer operations."),  # Common condition
                ("Transfers to offshore accounts require legal documentation.", 
                 "Documentation ensures compliance with regulations."),
                ("Recurring transfers can be paused for up to 30 days.", 
                 "Pause option provides flexibility for users."),
                ("Transfers under $100 are processed without fees.", 
                 "No fees incentivize small transaction volume."),
                ("Transfer delays are notified via email.", 
                 "Notifications keep customers informed of issues.")
            ],
            [
                ("All transfers must include a unique transaction ID.", 
                 "This ensures traceability for all transfer operations."),  # Common condition
                ("Transfers to business accounts require tax ID verification.", 
                 "Verification ensures proper business registration."),
                ("Emergency transfers are processed within 2 hours.", 
                 "Fast processing supports urgent needs."),
                ("Transfers between countries require currency conversion.", 
                 "Conversion rates are applied at the time of transfer."),
                ("Transfer fees are waived for first-time users.", 
                 "Waiver encourages new customer adoption.")
            ],
            # --- New version with cross-transaction relationships ---
            [
                ("Transfers that fail due to incorrect recipient details are refunded within 48 hours.",
                 "Refunds are processed to the original payment method."),
                ("Transfers made as payment for goods are refundable if the goods are returned and meet refund conditions.",
                 "Refunds for such transfers follow the standard refund process."),
                ("Transfers between linked accounts can be reversed within 24 hours upon request.",
                 "Reversals are expedited for customer convenience."),
                ("Transfers used for exchanges are adjusted based on the price difference.",
                 "If the exchanged item is cheaper, the difference is refunded; if more expensive, additional transfer is required."),
                ("International transfers are non-refundable once processed, except in case of failure.",
                 "Failed international transfers are refunded automatically."),
                ("Transfers for bulk orders require manager approval for refunds.",
                 "Approval ensures compliance with bulk purchase terms."),
                ("Transfers made by cryptocurrency are refundable at the current exchange rate if the transfer fails.",
                 "Conversion fees may apply for cryptocurrency refunds."),
                ("Transfers for promotional items are non-refundable.",
                 "Promotional item transfers are final."),
                ("Transfers for items received via exchange are subject to the exchange policy.",
                 "Refunds for exchanged items are only available if defective."),
                ("Transfers to unverified accounts are automatically refunded.",
                 "Unverified account transfers are not permitted.")
            ]
        ]
    },
    "exchanges": {
        "section_titles": ["Exchange Criteria", "Exchange Process", "Exchange Terms"],
        "versions": [
            # ...existing versions...
            [
                ("Exchanges are allowed within 7 days for defective items.", 
                 "Defective items qualify for exchange within this period."),
                ("Exchanged items must be in original, unused condition.", 
                 "Unused condition ensures items are suitable for resale."),
                ("No exchanges are available for international orders.", 
                 "International orders are excluded due to shipping complexities."),
                ("Exchange requests require a receipt or proof of purchase.", 
                 "Proof of purchase verifies the original transaction."),
                ("Additional fees may apply for exchanges due to price differences.", 
                 "Price differences are calculated at the time of exchange.")
            ],
            [
                ("Exchanges are permitted within 10 days for faulty products.", 
                 "Faulty products are eligible for exchange within this timeframe."),
                ("Items for exchange must be in original packaging and unused.", 
                 "Original packaging ensures items are in resalable condition."),
                ("International exchanges are subject to additional shipping fees.", 
                 "Shipping fees cover costs for international returns."),
                ("Proof of purchase is required for all exchange requests.", 
                 "A receipt or order confirmation is necessary for processing."),
                ("Price differences for exchanges may incur a 5% fee.", 
                 "This fee covers administrative costs for price adjustments.")
            ],
            [
                ("Exchanges are available within 5 days for defective goods.", 
                 "Defective goods must be reported promptly for exchange."),
                ("Exchanged items must be returned in saleable condition.", 
                 "Saleable condition means no damage or signs of use."),
                ("No exchanges for items purchased during promotional events.", 
                 "Promotional items are final sale due to discounts."),
                ("Exchange requests must include the original receipt.", 
                 "Receipts verify the purchase for exchange eligibility."),
                ("Additional costs for exchanges are calculated at checkout.", 
                 "Costs are determined based on the new item’s price.")
            ],
            [
                ("All exchanges require proof of original purchase.", 
                 "This ensures eligibility for all exchange requests."),  # Common condition
                ("Exchanges for oversized items require prior approval.", 
                 "Approval manages logistics for large items."),
                ("Same-day exchanges are free if within store hours.", 
                 "Free exchanges incentivize quick returns."),
                ("Exchanged items can be upgraded with a surcharge.", 
                 "Surcharge reflects the price difference for upgrades."),
                ("No exchanges are allowed after store closing time.", 
                 "Closing time marks the end of exchange eligibility.")
            ],
            [
                ("All exchanges require proof of original purchase.", 
                 "This ensures eligibility for all exchange requests."),  # Common condition
                ("Exchanges for digital products are not permitted.", 
                 "Digital items are non-returnable due to licensing."),
                ("Exchanges can be processed online with a tracking number.", 
                 "Tracking ensures secure return shipment."),
                ("Partial exchanges are allowed for multi-item orders.", 
                 "Partial returns adjust for unused items."),
                ("Exchanges are delayed during inventory audits.", 
                 "Audits ensure accurate stock levels.")
            ],
            [
                ("All exchanges require proof of original purchase.", 
                 "This ensures eligibility for all exchange requests."),  # Common condition
                ("Exchanges for discontinued items are limited to stock.", 
                 "Availability depends on remaining inventory."),
                ("Express exchanges are available for a 10% fee.", 
                 "Fee covers faster processing and shipping."),
                ("Exchanged items must be inspected within 24 hours.", 
                 "Inspection verifies condition before restocking."),
                ("No exchanges for items damaged by customer misuse.", 
                 "Misuse voids exchange eligibility.")
            ],
            [
                ("All exchanges require proof of original purchase.", 
                 "This ensures eligibility for all exchange requests."),  # Common condition
                ("Exchanges for limited-edition items are not allowed.", 
                 "Limited items are non-returnable due to rarity."),
                ("Exchanges can be requested via live chat support.", 
                 "Chat provides real-time assistance for exchanges."),
                ("Bulk exchanges require a 48-hour processing time.", 
                 "Processing ensures accurate handling of multiple items."),
                ("Exchanges are free if the item is out of stock.", 
                 "Free exchanges compensate for unavailability.")
            ],
            [
                ("All exchanges require proof of original purchase.", 
                 "This ensures eligibility for all exchange requests."),  # Common condition
                ("Exchanges for refurbished items have a 7-day limit.", 
                 "Limit applies due to warranty restrictions."),
                ("Exchanged items can be donated with a tax receipt.", 
                 "Receipts support charitable contributions."),
                ("Exchanges for wrong sizes are prioritized.", 
                 "Priority ensures quick resolution of sizing issues."),
                ("No exchanges during black Friday sales.", 
                 "Sales periods exclude exchange processing.")
            ],
            [
                ("All exchanges require proof of original purchase.", 
                 "This ensures eligibility for all exchange requests."),  # Common condition
                ("Exchanges for holiday-themed items end post-season.", 
                 "Post-season marks the end of return eligibility."),
                ("Exchanges include a free return label for first returns.", 
                 "Label encourages initial return attempts."),
                ("Exchanged items must match original purchase quantity.", 
                 "Quantity ensures fair exchange terms."),
                ("Exchanges are denied if serial numbers are removed.", 
                 "Serial numbers are required for verification.")
            ],
            # --- New version with cross-transaction relationships ---
            [
                ("Exchanges are permitted for items purchased by credit card, subject to refund and exchange policies.",
                 "Credit card exchanges may be refunded if the item is unavailable."),
                ("Exchanges for items paid by bank transfer require the original account for any refund.",
                 "Refunds for exchanges are processed within 7 business days."),
                ("Exchanges for items paid by cryptocurrency are subject to conversion fees for any refund.",
                 "Refunds are processed at the current exchange rate."),
                ("If an exchange cannot be fulfilled due to out-of-stock items, a refund will be issued.",
                 "Refunds for unfulfilled exchanges follow the standard refund process."),
                ("Exchanges for bulk orders require manager approval and may be limited.",
                 "Approval ensures compliance with bulk purchase and exchange terms."),
                ("Exchanges for promotional items are not eligible for refunds.",
                 "Promotional item exchanges are final."),
                ("Exchanges for items received via transfer are subject to transfer and exchange policies.",
                 "Refunds for such items are only available if defective."),
                ("Exchanges for failed transfers are not permitted; a refund will be issued instead.",
                 "Failed transfer exchanges are automatically refunded."),
                ("Exchanges between linked accounts are processed within 24 hours.",
                 "Linked account exchanges are expedited for convenience."),
                ("Exchanges for international orders may incur additional fees and have limited refund options.",
                 "Check international exchange terms for details.")
            ]
        ]
    }
}

def create_pdf_file(transaction_type: str, conditions: list, version: int, section_titles: list, output_dir: str) -> None:
    """Generate a PDF file with distributed transaction conditions and well-structured table."""
    filename = os.path.join(output_dir, f"{transaction_type}_v{version}.pdf")
    doc = SimpleDocTemplate(filename, pagesize=letter, leftMargin=0.5*inch, rightMargin=0.5*inch, topMargin=0.5*inch, bottomMargin=0.5*inch)
    styles = getSampleStyleSheet()
    story = []

    # Title
    story.append(Paragraph(f"{transaction_type.capitalize()} Policy - Version {version}", styles['Title']))
    story.append(Spacer(1, 12))

    # Purpose
    story.append(Paragraph("1. Purpose", styles['Heading2']))
    story.append(Paragraph(
        f"This document outlines the policy for {transaction_type} handled by our organization.",
        styles['Normal']
    ))
    story.append(Spacer(1, 12))

    # Policy Summary
    story.append(Paragraph("2. Policy Summary", styles['Heading2']))
    story.append(Paragraph(
        f"We aim to ensure all {transaction_type} are processed efficiently. "
        f"Contact support@company.com for issues or clarifications.",
        styles['Normal']
    ))
    story.append(Spacer(1, 12))

    # Distribute conditions across sections
    for i, (condition, explanation) in enumerate(conditions[:2], 3):
        story.append(Paragraph(f"{i}. {section_titles[0]}", styles['Heading2']))
        story.append(Paragraph(condition, styles['Normal']))
        story.append(Paragraph(explanation, styles['Normal']))
        story.append(Spacer(1, 12))

    story.append(Paragraph(f"{i+1}. {section_titles[1]}", styles['Heading2']))
    condition, explanation = conditions[2]
    story.append(Paragraph(condition, styles['Normal']))
    story.append(Paragraph(explanation, styles['Normal']))
    story.append(Spacer(1, 12))

    story.append(Paragraph(f"{i+2}. {section_titles[2]}", styles['Heading2']))
    for condition, explanation in conditions[3:]:
        story.append(Paragraph(condition, styles['Normal']))
        story.append(Paragraph(explanation, styles['Normal']))
        story.append(Spacer(1, 6))

    # Summary Table
    story.append(Paragraph(f"{i+3}. Summary", styles['Heading2']))
    styles['Normal'].fontSize = 9  # Set font size for table content
    table_data = [["Condition Type", "Details", "Explanation"]]
    for condition, explanation in conditions:
        table_data.append([Paragraph(transaction_type.capitalize(), styles['Normal']),
                          Paragraph(condition, styles['Normal']),
                          Paragraph(explanation, styles['Normal'])])
    colWidths = [70, 180, 320]  # Total: 570 points, fits within 612-point page with margins
    table = Table(table_data, colWidths=colWidths, repeatRows=1, rowHeights=None)
    table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 0), (-1, 0), 10),
        ('FONTSIZE', (0, 1), (-1, -1), 9),
        ('LEFTPADDING', (0, 0), (-1, -1), 10),
        ('RIGHTPADDING', (0, 0), (-1, -1), 10),
        ('TOPPADDING', (0, 0), (-1, -1), 10),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 10),
        ('BACKGROUND', (0, 1), (-1, -1), colors.white),
        ('GRID', (0, 0), (-1, -1), 1, colors.black)
    ]))
    story.append(table)
    story.append(Spacer(1, 12))

    # Contact
    story.append(Paragraph(f"{i+4}. Contact", styles['Heading2']))
    story.append(Paragraph(
        "For queries, email support@company.com or call 0800-123-456 during business hours (9 AM - 5 PM, Monday to Friday).",
        styles['Normal']
    ))

    doc.build(story)
    print(f"Created {filename}")

def create_docx_file(transaction_type: str, conditions: list, version: int, section_titles: list, output_dir: str) -> None:
    """Generate a DOCX file with distributed transaction conditions."""
    filename = os.path.join(output_dir, f"{transaction_type}_v{version}.docx")
    doc = Document()

    # Title
    doc.add_heading(f"{transaction_type.capitalize()} Policy - Version {version}", 0)

    # Purpose
    doc.add_heading("1. Purpose", level=2)
    doc.add_paragraph(
        f"This document outlines the policy for {transaction_type} handled by our organization."
    )

    # Policy Summary
    doc.add_heading("2. Policy Summary", level=2)
    doc.add_paragraph(
        f"We aim to ensure all {transaction_type} are processed efficiently. "
        f"Contact support@company.com for issues or clarifications."
    )

    # Distribute conditions across sections
    for i, (condition, explanation) in enumerate(conditions[:2], 3):
        doc.add_heading(f"{i}. {section_titles[0]}", level=2)
        doc.add_paragraph(condition)
        doc.add_paragraph(explanation)

    doc.add_heading(f"{i+1}. {section_titles[1]}", level=2)
    condition, explanation = conditions[2]
    doc.add_paragraph(condition)
    doc.add_paragraph(explanation)

    doc.add_heading(f"{i+2}. {section_titles[2]}", level=2)
    for condition, explanation in conditions[3:]:
        doc.add_paragraph(condition)
        doc.add_paragraph(explanation)

    # Summary Table
    doc.add_heading(f"{i+3}. Summary", level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Condition Type"
    hdr_cells[1].text = "Details"
    hdr_cells[2].text = "Explanation"
    for condition, explanation in conditions:
        row_cells = table.add_row().cells
        row_cells[0].text = transaction_type.capitalize()
        row_cells[1].text = condition
        row_cells[2].text = explanation

    # Contact
    doc.add_heading(f"{i+4}. Contact", level=2)
    doc.add_paragraph(
        "For queries, email support@company.com or call 0800-123-456 during business hours (9 AM - 5 PM, Monday to Friday)."
    )

    doc.save(filename)
    print(f"Created {filename}")

def create_xlsx_file(transaction_type: str, conditions: list, version: int, section_titles: list, output_dir: str) -> None:
    """Generate an XLSX file with distributed transaction conditions."""
    filename = os.path.join(output_dir, f"{transaction_type}_v{version}.xlsx")
    workbook = openpyxl.Workbook()
    sheet = workbook.active
    sheet.title = f"{transaction_type.capitalize()} Policy"

    # Headers
    sheet['A1'] = "Transaction Type"
    sheet['B1'] = "Section"
    sheet['C1'] = "Condition"
    sheet['D1'] = "Explanation"

    # Data
    for i, (condition, explanation) in enumerate(conditions[:2], 2):
        sheet[f'A{i}'] = transaction_type.capitalize()
        sheet[f'B{i}'] = section_titles[0]
        sheet[f'C{i}'] = condition
        sheet[f'D{i}'] = explanation

    i += 1
    sheet[f'A{i}'] = transaction_type.capitalize()
    sheet[f'B{i}'] = section_titles[1]
    condition, explanation = conditions[2]
    sheet[f'C{i}'] = condition
    sheet[f'D{i}'] = explanation

    for j, (condition, explanation) in enumerate(conditions[3:], i+1):
        sheet[f'A{j}'] = transaction_type.capitalize()
        sheet[f'B{j}'] = section_titles[2]
        sheet[f'C{j}'] = condition
        sheet[f'D{j}'] = explanation

    workbook.save(filename)
    print(f"Created {filename}")

def create_txt_file(transaction_type: str, conditions: list, version: int, section_titles: list, output_dir: str) -> None:
    """Generate a TXT file with distributed transaction conditions."""
    filename = os.path.join(output_dir, f"{transaction_type}_v{version}.txt")
    with open(filename, 'w', encoding='utf-8') as f:
        f.write(f"{transaction_type.capitalize()} Policy - Version {version}\n\n")
        f.write("1. Purpose\n")
        f.write(f"This document outlines the policy for {transaction_type} handled by our organization.\n\n")
        f.write("2. Policy Summary\n")
        f.write(f"We aim to ensure all {transaction_type} are processed efficiently. ")
        f.write("Contact support@company.com for issues or clarifications.\n\n")
        
        for i, (condition, explanation) in enumerate(conditions[:2], 3):
            f.write(f"{i}. {section_titles[0]}\n")
            f.write(f"{condition}\n")
            f.write(f"{explanation}\n\n")
        
        f.write(f"{i+1}. {section_titles[1]}\n")
        condition, explanation = conditions[2]
        f.write(f"{condition}\n")
        f.write(f"{explanation}\n\n")
        
        f.write(f"{i+2}. {section_titles[2]}\n")
        for condition, explanation in conditions[3:]:
            f.write(f"{condition}\n")
            f.write(f"{explanation}\n")
        
        f.write(f"\n{i+3}. Summary\n")
        for condition, explanation in conditions:
            f.write(f"- {transaction_type.capitalize()}: {condition} ({explanation})\n")
        
        f.write(f"\n{i+4}. Contact\n")
        f.write("For queries, email support@company.com or call 0800-123-456 during business hours (9 AM - 5 PM, Monday to Friday).\n")
    print(f"Created {filename}")

def generate_training_files(output_dir: str = "C:/Users/MSI/Desktop/CHATBOT/Flask/training_data") -> None:
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    for transaction_type, data in TRANSACTION_TYPES.items():
        section_titles = data["section_titles"]
        for version, conditions in enumerate(data["versions"], 1):
            create_pdf_file(transaction_type, conditions, version, section_titles, output_dir)
            create_docx_file(transaction_type, conditions, version, section_titles, output_dir)
            create_xlsx_file(transaction_type, conditions, version, section_titles, output_dir)
            create_txt_file(transaction_type, conditions, version, section_titles, output_dir)

if __name__ == "__main__":
    generate_training_files()